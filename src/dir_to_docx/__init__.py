import click
import os
import subprocess
import fnmatch

from typing import cast

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

from pygments import lex
from pygments.util import ClassNotFound
from pygments.lexers import guess_lexer_for_filename, get_lexer_by_name
from pygments.token import Token

# https://github.com/python-openxml/python-docx/issues/25#issuecomment-3211622482
def list_number(doc, par, prev=None, level=None, numbered=True):
    xpath_options = {
        True: {"single": "count(w:lvl)=1 and ", "level": 0},
        False: {"single": "", "level": level},
    }

    def style_xpath(prefer_single=True):
        style = par.style.style_id
        return (
            'w:abstractNum['
            '{single}w:lvl[@w:ilvl="{level}"]/w:pStyle[@w:val="{style}"]'
            ']/@w:abstractNumId'
        ).format(style=style, **xpath_options[prefer_single])

    def type_xpath(prefer_single=True):
        t = "decimal" if numbered else "bullet"
        return (
            'w:abstractNum['
            '{single}w:lvl[@w:ilvl="{level}"]/w:numFmt[@w:val="{type}"]'
            ']/@w:abstractNumId'
        ).format(type=t, **xpath_options[prefer_single])

    def get_abstract_id(numbering):
        for fn in (style_xpath, type_xpath):
            for prefer_single in (True, False):
                xp = fn(prefer_single)
                ids = numbering.xpath(xp)
                if ids:
                    return min(int(x) for x in ids)
        return 0

    if (prev is None or
        prev._p.pPr is None or
        prev._p.pPr.numPr is None or
        prev._p.pPr.numPr.numId is None):
        if level is None:
            level = 0
        numbering = doc.part.numbering_part.numbering_definitions._numbering
        anum = get_abstract_id(numbering)
        num = numbering.add_num(anum)
        num.add_lvlOverride(ilvl=level).add_startOverride(1)
        num_id = num.numId
    else:
        if level is None:
            level = prev._p.pPr.numPr.ilvl.val
        num_id = prev._p.pPr.numPr.numId.val

    ppr = par._p.get_or_add_pPr()
    numPr = ppr.get_or_add_numPr()
    numPr.get_or_add_numId().val = num_id
    numPr.get_or_add_ilvl().val = level

# https://github.com/python-openxml/python-docx/issues/359#issuecomment-2540346350
def pre_ref(run, cross_refs: list):
    for cref in cross_refs:
        bookmarkStart = OxmlElement("w:bookmarkStart")
        bookmarkStart.set(qn("w:id"), f"{cref['attrs']['w:id']}")
        bookmarkStart.set(qn("w:name"), cref["attrs"]["w:name"])
        run._r.append(bookmarkStart)
def post_ref(run, cross_refs: list):
    for cref in cross_refs:
        bookmarkEnd = OxmlElement("w:bookmarkEnd")
        bookmarkEnd.set(qn("w:id"), f"{cref['attrs']['w:id']}")
        run._r.append(bookmarkEnd)
def add_label(paragraph: Paragraph, label_type, refname: str, cross_refs: list, prefix: str):
    paragraph.add_run("")

    # numbering field
    run = paragraph.add_run(f"{prefix} ", style="Strong")
    pre_ref(run, cross_refs)

    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar)

    instrText = OxmlElement("w:instrText")
    instrText.text = f" SEQ {label_type} \\* ARABIC"
    run._r.append(instrText)

    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar)
    run.add_text(": ")
    post_ref(run, cross_refs)
def add_ref_place(paragraph, token):
    # caption type
    run = paragraph.add_run("")

    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar)

    ref_name = token["attrs"]["w:name"]
    instrText = OxmlElement("w:instrText")
    instrText.text = f" REF {ref_name} \\h"
    run._r.append(instrText)

    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar)


def guess_lexer_for_filename_(filename: str, content: str, associations: list[tuple[str, str]]):
    for association in associations:
        if fnmatch.fnmatch(filename, association[0]):
            return get_lexer_by_name(association[1])
    return guess_lexer_for_filename(filename, content)

@click.command("dir-to-docx")
@click.argument("input_dir", type=click.Path(exists=True, file_okay=False, dir_okay=True, path_type=str))
@click.argument("output_file", type=click.File("wb", encoding="utf-8", lazy=False))
@click.argument("path_prefix", type=click.Path(exists=True, file_okay=False, dir_okay=True, path_type=str))
@click.option("--exclude", type=str, multiple=True, help="many exclusions, e.g.: `*.ico`")
@click.option("--associate", type=tuple[str, str], multiple=True, help="many associations, e.g.: `*.svg,xml`")
def main(input_dir: str, output_file: click.File, path_prefix: str, exclude: tuple[str], associate: tuple[tuple[str, str]]):
    associations = list(associate)
    associations.append(("*.xml", "svg"))

    files = subprocess.check_output("git ls-files", shell=True, cwd=input_dir).decode("utf-8").splitlines()
    excludes = list(exclude)
    excludes.extend(["*.ico", "*.docx", "*.pptx", "*.xslx"])
    for e in excludes:
        files = list(filter(lambda f: not fnmatch.fnmatch(f, e), files))

    doc = Document()

    style = doc.styles.add_style("Code File", style_type=WD_STYLE_TYPE.TABLE)
    style.hidden = False
    style.quick_style = True
    style.locked = False
    style = doc.styles.add_style("String", style_type=WD_STYLE_TYPE.CHARACTER)
    style.hidden = False
    style.quick_style = True
    style.locked = False
    style = doc.styles.add_style("Number", style_type=WD_STYLE_TYPE.CHARACTER)
    style.hidden = False
    style.quick_style = True
    style.locked = False
    style = doc.styles.add_style("Comment", style_type=WD_STYLE_TYPE.CHARACTER)
    style.hidden = False
    style.quick_style = True
    style.locked = False
    style = doc.styles.add_style("Variable", style_type=WD_STYLE_TYPE.CHARACTER)
    style.hidden = False
    style.quick_style = True
    style.locked = False
    style = doc.styles.add_style("Keyword", style_type=WD_STYLE_TYPE.CHARACTER)
    style.hidden = False
    style.quick_style = True
    style.locked = False
    style = doc.styles.add_style("Operator", style_type=WD_STYLE_TYPE.CHARACTER)
    style.hidden = False
    style.quick_style = True
    style.locked = False

    cross_refs = []
    i = 1
    for file in files:
        path = os.path.relpath(os.path.join(input_dir, file), path_prefix).replace("\\", "/")
        cross_refs.append({"attrs": {"w:id": str(i), "w:name": path}})
        i += 1

    # print(len(cross_refs))

    i = 1
    with click.progressbar(files, show_pos=True, item_show_func=lambda f: os.path.relpath(os.path.join(input_dir, f), path_prefix).replace("\\", "/") if f is not None else None) as bar:
        for file in bar:
            path = os.path.relpath(os.path.join(input_dir, file), path_prefix).replace("\\", "/")
            # print(path, i)
            # caption
            caption = doc.add_paragraph()
            add_label(
                caption, label_type="File", refname=path, cross_refs=cross_refs, prefix="File "
            )
            caption.add_run(path)
            # content
            if file.endswith(".png") or file.endswith(".jpg") or file.endswith(".jpeg") or file.endswith(".gif"):
                image = doc.add_picture(os.path.join(input_dir, file))
            else:
                table = doc.add_table(rows=1, cols=1, style="Code File")
                with open(os.path.join(input_dir, file), "r", encoding="utf-8") as f:
                    content = f.read()
                    try:
                        tokens = lex(content,  guess_lexer_for_filename_(file, content, associations))
                        prev = None
                        paragraph = table.columns[0].cells[0].paragraphs[0]
                        paragraph.style = "List Number"
                        paragraph.add_run("\u200B")
                        for token in tokens:
                            if len(token[1]) == 0: continue
                            if token[0] in Token.Text.Whitespace:
                                if '\n' in token[1]:
                                    for line in token[1].split('\n')[1:]:
                                        list_number(doc, paragraph, prev=prev)
                                        prev = paragraph
                                        paragraph = table.columns[0].cells[0].add_paragraph("\u200B", style="List Number")
                                        paragraph.add_run(line)
                                else:
                                    paragraph.add_run(token[1])
                            elif token[0] in Token.Punctuation:
                                paragraph.add_run(token[1])
                            elif token[0] in Token.Error:
                                paragraph.add_run(token[1])
                            elif token[0] in Token.Comment:
                                paragraph.add_run(token[1], style="Comment")
                            elif token[0] in Token.Literal.String:
                                paragraph.add_run(token[1], style="String")
                            elif token[0] in Token.Literal.Number:
                                paragraph.add_run(token[1], style="Number")
                            elif token[0] in Token.Name:
                                paragraph.add_run(token[1], style="Variable")
                            elif token[0] in Token.Keyword:
                                paragraph.add_run(token[1], style="Keyword")
                            elif token[0] in Token.Operator:
                                paragraph.add_run(token[1], style="Operator")
                            else:
                                # print(token)
                                paragraph.add_run(token[1])
                    except ClassNotFound as err:
                        # print(err)
                        prev = None
                        for line in content.splitlines():
                            p = table.columns[0].cells[0].add_paragraph("\u200B" + line.rstrip(), style="List Number")
                            list_number(doc, p, prev=prev)
                            prev = p

        i += 1


    doc.save(output_file) # pyright: ignore[reportArgumentType]


# if __name__ == "__main__":
#     parser = argparse.ArgumentParser()
#     parser.add_argument("input_dir", type=str)
#     parser.add_argument("output_file", type=str)
#     parser.add_argument("path_prefix", type=str)
#     parser.add_argument("--exclude", type=str, nargs="*")
#     parser.add_argument("--associate", type=str, nargs="*", help="many associations, e.g.: `*.svg,xml`")
#     parser.add_help = True
#     args = parser.parse_args()

#     main(args.input_dir, args.output_file, args.path_prefix, args.exclude, list(map(lambda a: a.split(","), args.associate)))
